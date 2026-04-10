import io
import re
from pathlib import Path

try:
    import fitz
except ModuleNotFoundError:
    fitz = None

try:
    from docx import Document
except ModuleNotFoundError:
    Document = None

try:
    from cv_engine import sanitize_block, sanitize_line, sanitize_string_list
except ModuleNotFoundError:
    from backend.cv_engine import sanitize_block, sanitize_line, sanitize_string_list


MAX_UPLOAD_SIZE_BYTES = 8 * 1024 * 1024
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
SECTION_ALIASES = {
    "experience": ["experience", "experiences", "employment", "work", "internship", "internships", "stage", "alternance"],
    "education": ["education", "formation", "studies", "academic", "school", "university"],
    "projects": ["projects", "project", "portfolio", "selected work", "case studies", "projets", "realisations"],
    "skills": ["skills", "competences", "competencies", "stack", "tools", "technologies", "technical skills"],
    "languages": ["languages", "langues"],
    "certifications": ["certifications", "certification", "certificats"],
    "summary": ["profile", "summary", "about", "profil", "objective"],
}
SKILL_LIBRARY = [
    "Python", "SQL", "JavaScript", "TypeScript", "React", "Next.js", "Node.js", "FastAPI", "Django",
    "Flask", "HTML", "CSS", "Tailwind", "PostgreSQL", "MySQL", "MongoDB", "SQLite", "Pandas",
    "NumPy", "scikit-learn", "TensorFlow", "PyTorch", "Docker", "Git", "GitHub", "Power BI",
    "Tableau", "dbt", "Looker", "Excel", "Figma", "Notion", "Airtable", "Zapier", "n8n",
    "Make", "AWS", "Azure", "GCP", "Supabase", "Firebase", "REST API", "GraphQL",
]
EMAIL_RE = re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.I)
PHONE_RE = re.compile(r"(?:(?:\+\d{1,3}[\s.-]*)?(?:\(?\d{1,4}\)?[\s.-]*){2,}\d{2,4})")
URL_RE = re.compile(r"(https?://[^\s|]+|www\.[^\s|]+)", re.I)
DATE_RE = re.compile(r"\b(?:20\d{2}|19\d{2}|present|current|ongoing|aujourd'hui)\b", re.I)


def _extract_pdf_text(content: bytes) -> str:
    if fitz is None:
        raise RuntimeError("PyMuPDF is not installed")
    document = fitz.open(stream=content, filetype="pdf")
    try:
        return "\n\n".join(page.get_text("text") for page in document)
    finally:
        document.close()


def _extract_docx_text(content: bytes) -> str:
    if Document is None:
        raise RuntimeError("python-docx is not installed")
    document = Document(io.BytesIO(content))
    blocks = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                blocks.append(row_text)
    return "\n".join(blocks)


def extract_text_from_upload(filename: str, content_type: str, content: bytes) -> str:
    if not content:
        raise ValueError("Uploaded file is empty")
    if len(content) > MAX_UPLOAD_SIZE_BYTES:
        raise ValueError("Uploaded file is too large")

    extension = Path(filename or "").suffix.lower()
    normalized_type = (content_type or "").lower()
    raw_text = ""

    if extension and extension not in SUPPORTED_EXTENSIONS:
        raise ValueError("Unsupported file type. Use PDF, DOCX, TXT, or MD.")

    if extension == ".pdf" or normalized_type == "application/pdf":
        raw_text = _extract_pdf_text(content)
    elif extension == ".docx" or "wordprocessingml.document" in normalized_type:
        raw_text = _extract_docx_text(content)
    elif extension in {".txt", ".md"} or normalized_type.startswith("text/"):
        raw_text = content.decode("utf-8", errors="ignore")
    else:
        raise ValueError("Unsupported file type. Use PDF, DOCX, TXT, or MD.")

    cleaned = sanitize_block(raw_text, 32000)
    if len(cleaned.strip()) < 40:
        raise ValueError("The uploaded document did not contain enough readable text")
    return cleaned


def _normalize_heading(text: str) -> str:
    return re.sub(r"[^a-z0-9 ]+", " ", (text or "").lower()).strip()


def _non_empty_lines(text: str) -> list[str]:
    return [line.strip() for line in text.splitlines() if line.strip()]


def _split_blocks(text: str) -> list[str]:
    return [block.strip() for block in re.split(r"\n{2,}", text) if block.strip()]


def _find_heading_key(line: str) -> str | None:
    normalized = _normalize_heading(line)
    if len(normalized) > 40:
        return None
    for key, aliases in SECTION_ALIASES.items():
        if normalized in aliases:
            return key
        if any(normalized.startswith(alias) for alias in aliases):
            return key
    return None


def _looks_like_name(text: str) -> bool:
    clean = sanitize_line(text, 120)
    if not clean or any(token in clean.lower() for token in ["@", "http", "linkedin", "github", "curriculum", "resume", "cv"]):
        return False
    if re.search(r"\d", clean):
        return False
    words = clean.split()
    return 2 <= len(words) <= 4 and all(word[:1].isalpha() for word in words)


def _extract_contacts(text: str) -> dict:
    urls = URL_RE.findall(text)
    email = sanitize_line((EMAIL_RE.search(text) or [None])[0], 140) if EMAIL_RE.search(text) else ""
    phone = ""
    for match in PHONE_RE.findall(text):
        digits = re.sub(r"\D", "", match)
        if len(digits) >= 8:
            phone = sanitize_line(match, 60)
            break
    linkedin = ""
    github = ""
    website = ""
    for raw_url in urls:
        url = raw_url if raw_url.startswith("http") else f"https://{raw_url}"
        lowered = url.lower()
        if "linkedin.com" in lowered and not linkedin:
            linkedin = sanitize_line(url, 220)
        elif "github.com" in lowered and not github:
            github = sanitize_line(url, 220)
        elif not website:
            website = sanitize_line(url, 220)
    return {
        "email": email,
        "phone": phone,
        "linkedin": linkedin,
        "github": github,
        "website": website,
    }


def _extract_skills(text: str, max_items: int = 24) -> list[str]:
    normalized = _normalize_heading(text)
    found = []
    for skill in SKILL_LIBRARY:
        skill_norm = _normalize_heading(skill)
        if skill_norm and skill_norm in normalized and skill not in found:
            found.append(skill)
        if len(found) >= max_items:
            break
    return found


def _extract_list_section(text: str) -> list[str]:
    tokens = []
    for line in _non_empty_lines(text):
        clean = re.sub(r"^[\-\*\u2022]+\s*", "", line).strip()
        if _find_heading_key(clean):
            continue
        if len(clean) > 120 and "," not in clean:
            continue
        for token in re.split(r"[,\n|/]+", clean):
            entry = sanitize_line(token, 120)
            if entry:
                tokens.append(entry)
    return sanitize_string_list(tokens, max_items=20, max_length=120)


def _extract_section_map(text: str) -> dict[str, str]:
    sections = {key: [] for key in SECTION_ALIASES}
    current_key = "summary"
    for block in _split_blocks(text):
        lines = _non_empty_lines(block)
        if not lines:
            continue
        heading_key = _find_heading_key(lines[0])
        if heading_key:
            current_key = heading_key
            remainder = "\n".join(lines[1:]).strip()
            if remainder:
                sections[current_key].append(remainder)
            continue
        sections[current_key].append(block)
    return {key: "\n\n".join(value).strip() for key, value in sections.items() if value}


def _extract_dates(text: str) -> tuple[str, str]:
    matches = [sanitize_line(match.group(0), 40) for match in DATE_RE.finditer(text)]
    if len(matches) >= 2:
        return matches[0], matches[1]
    if len(matches) == 1:
        return matches[0], ""
    return "", ""


def _split_role_org(line: str) -> tuple[str, str]:
    clean = sanitize_line(line, 220)
    separators = [" at ", " @ ", " - ", " | ", " / "]
    for separator in separators:
        if separator in clean.lower():
            parts = re.split(re.escape(separator), clean, maxsplit=1, flags=re.I)
            if len(parts) == 2:
                left = sanitize_line(parts[0], 140)
                right = sanitize_line(parts[1], 140)
                return left, right
    return clean, ""


def _parse_experience_blocks(text: str) -> list[dict]:
    entries = []
    for block in _split_blocks(text):
        lines = _non_empty_lines(block)
        if not lines:
            continue
        start_date, end_date = _extract_dates(block)
        title, company = _split_role_org(lines[0])
        if len(lines) > 1 and not company and len(lines[1]) < 120:
            maybe_company = sanitize_line(lines[1], 140)
            if maybe_company and not EMAIL_RE.search(maybe_company):
                company = maybe_company
        summary_lines = [line for line in lines[1:] if line not in {company, start_date, end_date}]
        summary = sanitize_block(" ".join(summary_lines[:4]), 700)
        if title or company:
            entries.append(
                {
                    "company": company,
                    "title": title,
                    "location": "",
                    "start_date": start_date,
                    "end_date": end_date,
                    "summary": summary,
                    "highlights": sanitize_string_list(summary_lines[1:5], max_items=4, max_length=180),
                    "skills": _extract_skills(block, max_items=10),
                }
            )
        if len(entries) >= 8:
            break
    return entries


def _parse_education_blocks(text: str) -> list[dict]:
    entries = []
    for block in _split_blocks(text):
        lines = _non_empty_lines(block)
        if not lines:
            continue
        start_date, end_date = _extract_dates(block)
        degree, school = _split_role_org(lines[0])
        if len(lines) > 1 and not school:
            school = sanitize_line(lines[1], 140)
        summary = sanitize_block(" ".join(lines[1:4]), 500)
        entries.append(
            {
                "school": school,
                "degree": degree,
                "field": "",
                "location": "",
                "start_date": start_date,
                "end_date": end_date,
                "summary": summary,
                "highlights": sanitize_string_list(lines[2:5], max_items=3, max_length=180),
                "skills": _extract_skills(block, max_items=8),
            }
        )
        if len(entries) >= 6:
            break
    return [entry for entry in entries if entry["school"] or entry["degree"]]


def _parse_project_blocks(text: str) -> list[dict]:
    entries = []
    for block in _split_blocks(text):
        lines = _non_empty_lines(block)
        if not lines:
            continue
        urls = URL_RE.findall(block)
        summary = sanitize_block(" ".join(lines[1:4]), 700)
        entries.append(
            {
                "name": sanitize_line(lines[0], 140),
                "role": "",
                "url": sanitize_line(urls[0], 220) if urls else "",
                "summary": summary,
                "highlights": sanitize_string_list(lines[1:5], max_items=4, max_length=180),
                "technologies": _extract_skills(block, max_items=10),
            }
        )
        if len(entries) >= 8:
            break
    return [entry for entry in entries if entry["name"]]


def preparse_cv_text(text: str) -> dict:
    lines = _non_empty_lines(text)
    header_lines = lines[:8]
    contacts = _extract_contacts("\n".join(header_lines[:12] + lines[:40]))

    full_name = ""
    for line in header_lines[:4]:
        if _looks_like_name(line):
            full_name = sanitize_line(line, 120)
            break

    headline = ""
    for line in header_lines[1:6]:
        clean = sanitize_line(line, 140)
        if clean and clean not in contacts.values() and not _find_heading_key(clean):
            headline = clean
            break

    sections = _extract_section_map(text)
    summary_source = sections.get("summary") or "\n".join(header_lines[1:4])
    summary = sanitize_block(summary_source, 700)
    skills = sanitize_string_list(
        [*(_extract_list_section(sections.get("skills", ""))), *(_extract_skills(text, max_items=24))],
        max_items=30,
        max_length=80,
    )

    parsed = {
        "full_name": full_name,
        "headline": headline,
        "email": contacts["email"],
        "phone": contacts["phone"],
        "location": "",
        "website": contacts["website"],
        "linkedin": contacts["linkedin"],
        "github": contacts["github"],
        "summary": summary,
        "target_roles": sanitize_string_list([headline], max_items=4, max_length=80) if headline else [],
        "skills": skills,
        "languages": _extract_list_section(sections.get("languages", "")),
        "certifications": _extract_list_section(sections.get("certifications", "")),
        "education": _parse_education_blocks(sections.get("education", "")),
        "experience": _parse_experience_blocks(sections.get("experience", "")),
        "projects": _parse_project_blocks(sections.get("projects", "")),
        "raw_sections": {
            "summary": sanitize_block(sections.get("summary", ""), 1200),
            "skills": sanitize_block(sections.get("skills", ""), 1200),
            "experience": sanitize_block(sections.get("experience", ""), 3000),
            "education": sanitize_block(sections.get("education", ""), 2000),
            "projects": sanitize_block(sections.get("projects", ""), 2000),
        },
        "top_excerpt": sanitize_block("\n".join(lines[:40]), 2500),
        "notes": [
            note
            for note in [
                "Local pre-parser extracted the CV into sections before LLM normalization.",
                "LLM should refine this parsed structure, not re-read the full document from scratch.",
            ]
            if note
        ],
    }
    return parsed
